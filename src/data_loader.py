from pathlib import Path
import json
import zipfile

import pandas as pd
from docx import Document
from tqdm.auto import tqdm


def safe_str(value):
    if value is None:
        return ""
    return str(value).replace("\n", " ").replace("\t", " ").strip()


def unzip_competition_bundle(raw_dir):
    raw_dir = Path(raw_dir)
    zip_files = list(raw_dir.glob("*.zip"))

    if not zip_files:
        raise FileNotFoundError(f"No ZIP file found inside {raw_dir}")

    zip_path = zip_files[0]

    with zipfile.ZipFile(zip_path, "r") as z:
        z.extractall(raw_dir)

    return zip_path


def find_clean_file(raw_dir, filename):
    raw_dir = Path(raw_dir)

    matches = [
        p for p in raw_dir.rglob(filename)
        if "__MACOSX" not in str(p) and not p.name.startswith("._")
    ]

    if not matches:
        raise FileNotFoundError(f"{filename} not found inside {raw_dir}")

    return matches[0]


def get_bundle_paths(raw_dir):
    raw_dir = Path(raw_dir)

    unzip_competition_bundle(raw_dir)

    return {
        "candidates": find_clean_file(raw_dir, "candidates.jsonl"),
        "job_description": find_clean_file(raw_dir, "job_description.docx"),
        "validator": find_clean_file(raw_dir, "validate_submission.py"),
    }


def read_docx(path):
    doc = Document(path)
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def build_candidate_text(profile, career, education, skills, certifications, languages):
    title = safe_str(profile.get("current_title"))
    headline = safe_str(profile.get("headline"))
    summary = safe_str(profile.get("summary"))
    location = safe_str(profile.get("location"))
    country = safe_str(profile.get("country"))
    industry = safe_str(profile.get("current_industry"))
    company = safe_str(profile.get("current_company"))
    years = profile.get("years_of_experience", 0)

    skill_names = [safe_str(s.get("name")) for s in skills]

    skill_text = " ".join(
        f"{safe_str(s.get('name'))} {safe_str(s.get('proficiency'))} "
        f"{s.get('duration_months', 0)} months endorsements {s.get('endorsements', 0)}"
        for s in skills
    )

    career_text = " ".join(
        f"{safe_str(job.get('title'))} at {safe_str(job.get('company'))} "
        f"in {safe_str(job.get('industry'))}. {safe_str(job.get('description'))} "
        f"duration {job.get('duration_months', 0)} months current {job.get('is_current', False)}."
        for job in career
    )

    education_text = " ".join(
        f"{safe_str(e.get('degree'))} {safe_str(e.get('field_of_study'))} "
        f"{safe_str(e.get('institution'))} {safe_str(e.get('tier'))}"
        for e in education
    )

    cert_text = " ".join(safe_str(c) for c in certifications)

    language_text = " ".join(
        f"{safe_str(l.get('language'))} {safe_str(l.get('proficiency'))}"
        for l in languages
    )

    text = f"""
    Current role: {title}
    Headline: {headline}
    Summary: {summary}
    Skills: {" ".join(skill_names)}
    Skill evidence: {skill_text}
    Career history: {career_text}
    Education: {education_text}
    Certifications: {cert_text}
    Languages: {language_text}
    Location: {location}, {country}
    Experience: {years} years
    Industry: {industry}
    Company: {company}
    """

    return " ".join(text.lower().split())


def load_candidates(candidate_path):
    rows = []

    with open(candidate_path, "r", encoding="utf-8") as f:
        for line in tqdm(f, desc="Loading candidates"):
            if not line.strip():
                continue

            candidate = json.loads(line)

            profile = candidate.get("profile", {})
            career = candidate.get("career_history", [])
            education = candidate.get("education", [])
            skills = candidate.get("skills", [])
            certifications = candidate.get("certifications", [])
            languages = candidate.get("languages", [])
            signals = candidate.get("redrob_signals", {})

            semantic_text = build_candidate_text(
                profile=profile,
                career=career,
                education=education,
                skills=skills,
                certifications=certifications,
                languages=languages,
            )

            rows.append({
                "candidate_id": candidate.get("candidate_id"),
                "title": safe_str(profile.get("current_title")),
                "headline": safe_str(profile.get("headline")),
                "summary": safe_str(profile.get("summary")),
                "location": safe_str(profile.get("location")),
                "country": safe_str(profile.get("country")),
                "industry": safe_str(profile.get("current_industry")),
                "company": safe_str(profile.get("current_company")),
                "years_experience": float(profile.get("years_of_experience", 0) or 0),
                "skills": skills,
                "career_history": career,
                "education": education,
                "signals": signals,
                "semantic_text": semantic_text,
            })

    return pd.DataFrame(rows)